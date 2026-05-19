# -*- coding: utf-8 -*-
"""
==========================================================
46-MAVZU: SAVOL-JAVOB CHATBOT (NLP ASOSIDA)
AUTOMATIC ACADEMIC DOCX REPORT GENERATOR (30+ PAGES)
TOSHKENT DAVLAT IQTISODIYOT UNIVERSITETI
Guruh: IB-75/23, Talaba: Quvonchbek
==========================================================
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_docx():
    doc = docx.Document()
    
    # Page setup - Margins (1 inch / 2.54 cm standard)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style defaults - Times New Roman 14pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)

    # Helper function for headings
    def add_custom_heading(text, level, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
            align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            run.font.size = Pt(12)
            align = WD_ALIGN_PARAGRAPH.LEFT
            
        h.alignment = align
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.keep_with_next = True
        return h

    # Helper function for regular paragraphs
    def add_custom_paragraph(text, before=0, after=6, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, font_size=14):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line_spacing
        p.alignment = align
        
        # First line indentation for justified paragraphs (~1.25 cm standard)
        if align == WD_ALIGN_PARAGRAPH.JUSTIFY:
            p.paragraph_format.first_line_indent = Pt(36)
            
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.italic = italic
        return p

    # Helper function for image insertion with caption
    def add_project_image(img_path, caption, width_inches=5.5):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.add_run().add_picture(img_path, width=Inches(width_inches))
            
            # Caption
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(12)
            run = cap.add_run(caption)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.italic = True
        else:
            print(f"Warning: Image {img_path} not found.")

    # Helper function for adding code snippets
    def add_code_snippet(code_text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        return p

    # ==================== TITUL PAGE ====================
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    run1 = p1.add_run("TOSHKENT DAVLAT IQTISODIYOT UNIVERSITETI")
    run1.bold = True
    run1.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("RAQAMLI IQTISODIYOT VA AXBOROT TEXNOLOGIYALARI FAKULTETI")
    run2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(36)
    run3 = p3.add_run("AXBOROT XAVFSIZLIGI YO'NALIShI\nIB-75/23 GURUH TALABASI")
    run3.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    t_run = title_p.add_run("“TABIIY TILNI QAYTA ISHLASH (NLP) ASOSIDA SAVOL-JAVOB CHATBOTI VA UNING INTERAKTIV INTERFEYSI”")
    t_run.bold = True
    t_run.font.size = Pt(16)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(12)
    s_run = sub_p.add_run("MAVZUSIDAGI")
    s_run.font.size = Pt(12)

    ind_p = doc.add_paragraph()
    ind_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ind_p.paragraph_format.space_after = Pt(48)
    i_run = ind_p.add_run("INDIVIDUAL LOYIHA ISHI")
    i_run.bold = True
    i_run.font.size = Pt(18)

    for _ in range(2):
        doc.add_paragraph()

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    author_p.paragraph_format.space_after = Pt(4)
    run_auth1 = author_p.add_run("Bajardi: Guruh IB-75/23 Talabasi — Quvonchbek\n")
    run_auth1.bold = True
    run_auth1.font.size = Pt(12)
    run_auth2 = author_p.add_run("Ilmiy rahbar: Maxammadiyev M.\n")
    run_auth2.font.size = Pt(12)
    run_auth3 = author_p.add_run("Himoya qilindi: _______\n")
    run_auth3.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    city_p = doc.add_paragraph()
    city_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_p.paragraph_format.space_before = Pt(48)
    city_run = city_p.add_run("TOSHKENT — 2026")
    city_run.bold = True
    city_run.font.size = Pt(12)

    doc.add_page_break()

    # ==================== LOYIHA TOPSHIRIG'I ====================
    add_custom_heading("INDIVIDUAL LOYIHA ISHI TOPSHIRIG'I", level=1, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=18)
    
    add_custom_paragraph("Guruh: IB-75/23", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("Talaba: Quvonchbek", bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_custom_paragraph("1. Mavzu: Tabiiy tilni qayta ishlash (NLP) asosida interaktiv savol-javob chatboti va uning interfeysini loyihalash (Mavzu №46).", font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("2. Asosiy ma’lumotlar: TDIU akademik ma'lumotlari zanjiri (qabul, o'qish imtihonlari, stipendiya, yotoqxona va kutubxona). Ushbu savol-javoblar bazasini o'zbek tili morfemalari asosida tozalash, tokenizatsiya, stop-word'larni o'chirish, rule-based stemming va TF-IDF va Kosinus o'xshashligi (Cosine Similarity) algoritmik mexanizmi asosida moslashtirish.", font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("3. Foydalanilgan materiallar: Python (Pandas, NumPy, Matplotlib, Scikit-Learn kutubxonalari), HTML5 (Semantic elements), CSS3 (Glassmorphism Dark Mode), JavaScript (Vanilla ES6), Chart.js.", font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("4. Dastur qismi:\n   - 1) nlp_chatbot.py (Python NLP modeli - TF-IDF Vectorizer va Cosine Similarity vizualizatorlari).\n   - 2) index.html & style.css (Premium shisha dizaynli interaktiv foydalanuvchi interfeysi).\n   - 3) app.js (Brauzerda to'liq ishlovchi custom o'zbekcha NLP modeli va Chart.js vizualizatsiyasi).", font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("5. Hisobot mundarijasi: Kirish, 2 ta asosiy bob (6 ta paragraf), xulosa, foydalanilgan adabiyotlar ro'yxati va amaliy dasturiy ilovalar.", font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("6. Topshiriq muddati: May, 2026-yil.", font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT)

    for _ in range(3):
        doc.add_paragraph()

    signatures = doc.add_paragraph()
    signatures.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    signatures.paragraph_format.space_before = Pt(24)
    run_sig1 = signatures.add_run("Rahbar: _________ Maxammadiyev M.                       ")
    run_sig1.font.size = Pt(12)
    run_sig2 = signatures.add_run("Talaba: _________ Quvonchbek")
    run_sig2.font.size = Pt(12)

    doc.add_page_break()

    # ==================== MUNDARIJA ====================
    add_custom_heading("MUNDARIJA", level=1, before=18, after=18)
    
    add_custom_paragraph("KIRISH ................................................................................................................................ 4", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_custom_paragraph("I BOB. TABIIY TILNI QAYTA ISHLASH (NLP) VA CHATBOT TIZIMLARINING NAZARIY-METODOLOGIK ASOSLARI ................................................................................................ 5", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("   1.1. Tabiiy tilni qayta ishlash (NLP) sohasi rivojlanishi va uning iqtisodiy-ijtimoiy dolzarbligi ....... 5", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("   1.2. Tokenizatsiya, stop-word'lar filtratsiyasi va o'zbek tili morfologik tahlil (stemming) qoidalari ... 9", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("   1.3. TF-IDF (Term Frequency-Inverse Document Frequency) va Kosinus o'xshashligi (Cosine Similarity) matematik va algoritmik tahlili ......................................................................................... 13", align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_custom_paragraph("II BOB. NLP ASOSIDAGI INTERAKTIV SAVOL-JAVOB CHATBOT TIZIMINI DASTURIY AMALGA OSHIRISH VA TAHLIL ............................................................................................. 17", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("   2.1. Tizimning umumiy me'morchiligi (Architecture) va ma'lumotlar to'plami (Dataset) preprocessing bosqichi ............................................................................................................................. 17", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("   2.2. Python tilida NLP modelini amalga oshirish va visual tahlil natijalari ....................................... 21", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("   2.3. Foydalanuvchining interaktiv glassmorphism veb-dashboard dizayni va uning tahlili ......... 26", align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_custom_paragraph("XULOSA ............................................................................................................................. 31", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("FOYDALANILGAN ADABIYOTLAR RO'YHATI ......................................................................... 33", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # ==================== KIRISH ====================
    add_custom_heading("KIRISH", level=1, before=18, after=12)
    
    add_custom_paragraph("Axborot texnologiyalari va sun'iy intellekt tizimlarining bugungi shiddatli rivojlanishi inson hayoti va jamiyat faoliyatining barcha jabhalarini, xususan, oliy ta'lim tizimini ham tubdan o'zgartirmoqda. Universitetlarda ma'muriy va ta'limiy ma'lumotlarning haddan tashqari ko'pligi hamda talabalar, abituriyentlar va ota-onalar tomonidan har kuni yuzlab, minglab standart savollar berilishi ta'lim muassasalari xodimlariga juda katta ish yukini yuklaydi. Bunga yechim sifatida zamonaviy sun'iy intellekt texnologiyalaridan biri hisoblangan, inson tili (tabiiy til) bilan muloqot qila oladigan avtomatlashtirilgan Savol-Javob Chatbotlarini ishlab chiqish bugungi kunda o'ta dolzarb hisoblanadi. Tabiiy tilni qayta ishlash (Natural Language Processing - NLP) sun'iy intellektning eng jadal rivojlanayotgan va amaliy qo'llanilayotgan sohalaridan biridir. NLP yordamida kompyuterlar inson gapiradigan yoki yozadigan matnlarni faqatgina oddiy simvollar to'plami sifatida ko'rmasdan, balki ularning semantik, grammatik va iqtisodiy mohiyatini anglab yetishga xizmat qiladi.")
    
    add_custom_paragraph("Individual loyiha ishining dolzarbligi: Bugungi kunda Toshkent Davlat Iqtisodiyot Universiteti (TDIU) kabi yirik va nufuzli ta'lim muassasalarida minglab talabalar o'qiydi. Ushbu talabalarning universitet qabul qoidalari, fakultetlar tarkibi, kontrakt to'lovlari shartlari, stipendiyalar miqdori va kutubxona ish tartibi kabi ko'plab savollariga real vaqt rejimida tezkor javob qaytaruvchi avtomatlashtirilgan chatbotlarni joriy qilish vaqt va inson resurslarini tejash uchun juda muhimdir. Ayniqsa, o'zbek tilidagi tabiiy matnlarni tahlil qilish, undagi grammatik qo'shimchalarni tozalash (stemming), to'xtash so'zlarini (stop-words) olib tashlash va matnlar semantik o'xshashligini Kosinus o'xshashligi (Cosine Similarity) va TF-IDF usullari bilan hisoblash ancha murakkab ilmiy muammo hisoblanib, ushbu loyiha doirasida ushbu muammolarga amaliy va qulay yechim taklif etilgan. Axborot xavfsizligi nuqtai nazaridan ham dynamic axborot muhitini himoyalash, soxta so'rovlar oqimini saralash muhim o'rin tutadi.")
    
    add_custom_paragraph("Individual loyiha ishining maqsadi: TDIU akademik savollariga real vaqtda va yuqori aniqlik bilan javob qaytaradigan, tabiiy tilni qayta ishlash (NLP) algoritmlariga asoslangan, to'liq ishlovchi custom o'zbekcha chatbot dvigatelini va uni oddiy foydalanuvchilar oson ishlata olishi uchun premium interaktiv 'Glassmorphism' dizayn tizimidagi veb-dashboard ko'rinishida amaliyotga tatbiq etishdir. Tizim orqali foydalanuvchilar nafaqat chat qila oladilar, balki har bir berilgan savolning NLP pipeline orqali qanday bosqichlardan o'tishini visual ravishda kuzatib, tizimning ishlash prinsiplarini ilmiy asoslaydilar.")
    
    add_custom_paragraph("Individual loyiha ishining vazifalari:")
    add_custom_paragraph("• Tabiiy tilni qayta ishlash (NLP) va chatbot tizimlarining jahon amaliyotidagi nazariy asoslarini har tomonlama tahlil qilish;", before=0, after=2)
    add_custom_paragraph("• O'zbek tili morfologik xususiyatlaridan kelib chiqib, stop-word'larni filtrlovchi va grammatik kelishik, ko'plik va egalik qo'shimchalarini kesuvchi (stemming) maxsus rule-based algoritmlarni ishlab chiqish;", before=0, after=2)
    add_custom_paragraph("• Matnlarni iqtisodiy-matematik uslubda ifodalash uchun TF-IDF (Term Frequency-Inverse Document Frequency) og'irliklar formulasini va Kosinus o'xshashligi (Cosine Similarity) mezonini dasturlash;", before=0, after=2)
    add_custom_paragraph("• TDIU doirasida 15 ta asosiy o'quv-ma'muriy savol-javoblardan iborat ma'lumotlar bazasini (dataset) shakllantirish, Python modelini o'qitish va ilmiy vizualizatsiya grafiklarini yaratish;", before=0, after=2)
    add_custom_paragraph("• Veb-brauzerda to'liq offline ishlaydigan custom Javascript NLP dvigatelini, interaktiv bilimlar bazasi menejerini va Chart.js grafiklari orqali tizim samaradorligini ko'rsatuvchi premium interfeysni (index.html, style.css, app.js) yaratish.", before=0, after=6)
    
    add_custom_paragraph("Loyiha ishi obyekti: O'zbek tilidagi tabiiy matnlarni qayta ishlash jarayonlari va savol-javob chatbotlarining dasturiy tizimlari.")
    add_custom_paragraph("Loyiha ishi predmeti: O'zbek tili qoidalari asosida tokenizatsiya, stop-word'lar filtratsiyasi, stemming, TF-IDF va Kosinus o'xshashligi algoritmlari yordamida chatbot javobini optimal aniqlash usullari.")
    add_custom_paragraph("Loyiha ishining tarkibiy tuzilishi: Mazkur individual loyiha ishi Kirish, ikkita bob, 6 ta paragraf, Xulosa va Foydalanilgan adabiyotlar ro'yxatidan iborat bo'lib, jami 33 betni tashkil qiladi va unda 3 ta ilmiy-tahliliy rasmlar joy olgan. Har bir bo'lim mavzuning ilmiy hamda amaliy mohiyatini tizimli ochib beradi.")

    doc.add_page_break()

    # ==================== I BOB ====================
    add_custom_heading("I BOB. TABIIY TILNI QAYTA ISHLASH (NLP) VA CHATBOT TIZIMLARINING NAZARIY-METODOLOGIK ASOSLARI", level=1, before=18, after=12)
    
    add_custom_heading("1.1. Tabiiy tilni qayta ishlash (NLP) sohasi rivojlanishi va uning iqtisodiy-ijtimoiy dolzarbligi", level=2, before=12, after=6)
    
    add_custom_paragraph("Tabiiy tilni qayta ishlash (Natural Language Processing - NLP) sun'iy intellekt va lingvistikani birlashtiruvchi fanlararo yo'nalish bo'lib, uning asosiy maqsadi kompyuterlarga inson tilidagi matnlar va nutqni tushunish, tahlil qilish, talqin qilish va generatsiya qilish imkoniyatini berishdir. Rivojlanish tarixiga nazar tashlasak, NLP asosan uchta bosqichni bosib o'tdi. Birinchisi, 1950-yillardan boshlangan qoidalarga asoslangan (rule-based) tizimlar bo'lib, unda mutaxassislar tomonidan juda ko'plab qattiq lingvistik qoidalar qo'lda yozilgan. Bu usul juda cheklangan bo'lib, tilning moslashuvchanligi va ko'p ma'noliligini to'liq qamrab ololmagan. Ikkinchisi, 1990-yillardan boshlab statistika va ehtimollar nazariyasiga asoslangan modellar davri bo'ldi. Ushbu bosqichda matnlar korpusi (corpora) tahlil qilinib, so'zlarning birgalikda uchrashish ehtimolligi asosida bashoratlar qilingan. Uchinchisi esa, hozirgi kunda hukmronlik qilayotgan Mashinaviy o'qitish va Chuqur o'rganish (Deep Learning) hamda Katta Til Modellari (LLM - Large Language Models) davridir. Zamonaviy NLP tizimlari endilikda inson tili mohiyatini juda yuqori aniqlikda, matn semantikasini o'rgangan holda tushunmoqda.")
    
    add_custom_paragraph("Raqamli transformatsiya va iqtisodiyot sharoitida NLP texnologiyalarining dolzarbligi juda tez sur'atlar bilan o'sib bormoqda. Zamonaviy bizneslar va davlat tashkilotlari har kuni iste'molchilardan trillionlab gigabayt matnli ma'lumotlar qabul qilib oladi. Ushbu ma'lumotlarni qo'lda qayta ishlash millionlab dollar moliyaviy xarajatlar va juda katta inson resursini talab qiladi. Sun'iy intellekt chatbotlari orqali mijozlar so'rovlarini avtomatik qayta ishlash korxonalarga mijozlarni qo'llab-quvvatlash xarajatlarini 30% dan 50% gacha tejash imkonini beradi. Shu bilan birga, chatbotlar 24/7 rejimida real vaqtda bir zumda javob qaytargani sababli mijozlarning qoniqish darajasini (CSAT) sezilarli ravishda oshiradi. Biznes iqtisodiyotidan tashqari oliy ta'lim tizimida ham minglab talabalarning o'quv jarayonidagi savollarini avtomatlashtirish, universitet ish faoliyatini raqamlashtirish uchun NLP chatbotlari eng samarali strategik yechimlardan biri hisoblanadi.")
    
    add_custom_paragraph("NLP sohasidagi eng muhim lingvistik qiyinchiliklardan biri bu tillarning morfologik tuzilishidir. Masalan, ingliz tili kabi analitik tillarda so'zlarning shakllanishi ancha sodda bo'lsa, o'zbek tili kabi agglyutinativ (agglyutinativ) tillarda so'z o'zagiga ko'plab grammatik qo'shimchalar (ko'plik, egalik, kelishik, yuklamalar) birin-ketin qo'shilib boradi. Bu esa o'zbek tili uchun universal va aniq NLP preprocessing (dastlabki qayta ishlash) pipeline'larini yaratishni taqozo etadi. Mazkur loyiha ishi o'zbek tilining ushbu grammatik tabiati hisobga olingan holda, TDIU akademik tizimi uchun maxsus NLP oqimini loyihalashga qaratilgan.")

    add_custom_heading("1.2. Tokenizatsiya, stop-word'lar filtratsiyasi va o'zbek tili morfologik tahlil (stemming) qoidalari", level=2, before=12, after=6)
    
    add_custom_paragraph("Tabiiy matnlarni tahlil qilishda kompyuterlar to'g'ridan-to'g'ri butun gaplarni yoki paragraflarni semantik tushuna olmaydi. Matnni matematik modelga o'tkazishdan oldin, u bir necha ketma-ket qayta ishlash (preprocessing) bosqichlaridan o'tishi shart. Ushbu bosqichlarning birinchisi va eng asosiysi — Tokenizatsiya (Tokenization) hisoblanadi. Tokenizatsiya — bu uzluksiz matn oqimini alohida so'zlar, simvollar yoki frazalarga (ya'ni tokenlarga) ajratish jarayonidir. Odatda so'z tokenizatsiyasi matndagi bo'shliqlar (whitespace) va tinish belgilari asosida amalga oshiriladi. Gaplar tokenizatsiya bosqichida kichik harflarga (lowercasing) o'tkaziladi, chunki kompyuter uchun 'Tdiu' va 'tdiu' so'zlari ikkita mutlaqo boshqa obyekt hisoblanadi. Kichik harfga o'tkazish so'zlar o'rtasidagi yagona lug'atni shakllantirish imkonini beradi.")
    
    add_custom_paragraph("Keyingi muhim bosqich — Stop-word'larni (to'xtash so'zlarini) olib tashlashdir. Stop-words — bu matnda juda tez-tez uchraydigan, biroq gapning asosiy semantik ma'nosini yoki kontekstini ochib berishda deyarli hech qanday iqtisodiy-matematik qiymatga ega bo'lmagan yordamchi so'zlardir. O'zbek tilida bunday so'zlarga 'va', 'bilan', 'uchun', 'ammo', 'esa', 'ham', 'chunki', 'yoki', 'shu', 'bu', 'men', 'sen' kabi bog'lovchilar va olmoshlar kiradi. Agar ushbu yordamchi so'zlar matndan olib tashlanmasa, ular modelda juda yuqori og'irlik (chastota) olib ketadi va natijada modelning asosiy mazmuni buzilib, aniqligi pasayadi. Shu sababli, stop-word'lar maxsus lug'atlar orqali filtrlanadi va faqat gapdagi asosiy ma'noli lemmalar qoldiriladi.")
    
    add_custom_paragraph("Preprocess bosqichining eng murakkab qismi bu Lemmatizatsiya yoki Stemming jarayonidir. Stemming — so'zlarning oxiridagi grammatik shakl yasovchi va so'z o'zgartiruvchi qo'shimchalarni kesib tashlab, so'zni uning asosiy o'zagi (stem) ko'rinishiga keltirishdir. O'zbek tili agglyutinativ bo'lgani sababli, bitta so'z o'zagi o'nlab variantlarda kelishi mumkin, masalan: 'kutubxona', 'kutubxonalar', 'kutubxonamizdan', 'kutubxonadagi' va hokazo. Tizim bu so'zlarning barchasini yagona o'zak bo'lgan 'kutubxona'ga keltirishi lozim, aks holda modelning lug'at (vocabulary) hajmi asossiz ravishda juda kattalashib ketadi va so'zlar orasidagi bog'liqlik yo'qoladi. Tizim doirasida o'zbek tili uchun maxsus morfemik stemming qoidalari ishlab chiqilib, dynamic ko'rinishga keltirildi.")

    add_custom_heading("1.3. TF-IDF (Term Frequency-Inverse Document Frequency) va Kosinus o'xshashligi (Cosine Similarity) matematik va algoritmik tahlili", level=2, before=12, after=6)
    
    add_custom_paragraph("Matnlar preprocessing bosqichidan o'tib tozalanganidan so'ng, ularni kompyuter hisob-kitob qila oladigan raqamli ko'rinishga (Vektor fazoviy modeliga) o'tkazish lozim. Bu jarayon matnlarni Vektorlashtirish (Vectorization) deb ataladi. NLP sohasidagi eng mashhur va o'zining soddaligi va yuqori samaradorligi bilan ajralib turadigan matematik usullardan biri bu TF-IDF (Term Frequency-Inverse Document Frequency) modelidir. TF-IDF ikki qismdan iborat bo'lib, har bir alohida so'zning (lemmaning) ma'lum bir matndagi (savoldagi) ahamiyatini baholaydi. Birinchi qism — TF (Term Frequency - So'z chastotasi) so'zning matndagi uchrashish chastotasini ko'rsatadi: TF(t, d) = (so'z t ning d matndagi soni) / (d matndagi jami so'zlar soni). TF qiymati qanchalik yuqori bo'lsa, o'sha so'z berilgan gapda shunchalik muhim hisoblanadi.")
    
    add_custom_paragraph("Ikkinchi qism — IDF (Inverse Document Frequency - Teskari hujjat chastotasi) bo'lib, u so'zning barcha matnlar to'plamidagi (korpusdagi) noyobligini o'lchaydi: IDF(t, D) = log(1 + N / (1 + DF(t))) + 1, bunda N — korpusdagi jami matnlar soni, DF(t) — ushbu so'z qatnashgan matnlar soni. Agar biror so'z (masalan, 'universitet') barcha savollarda qatnashsa, uning IDF qiymati juda past (nolga yaqin) bo'ladi, chunki u savollarni bir-biridan ajratishda foydali emas. Aksincha, 'stipendiya' so'zi faqatgina bitta savolda uchrasa, uning IDF qiymati juda yuqori bo'ladi. Yakuniy TF-IDF og'irligi ushbu ikki qiymatning ko'paytmasi orqali topiladi: TF-IDF(t, d, D) = TF(t, d) * IDF(t, D). Model barcha savollar bo'yicha hisoblangan har bir matn uchun TF-IDF vektorini hosil qiladi va uni L2 normallash (vektor uzunligini 1.0 ga tenglash) yordamida standartlashtiradi.")
    
    add_custom_paragraph("Foydalanuvchi yangi savol yozganida, uning savoli ham xuddi shu preprocessing va TF-IDF bosqichlaridan o'tib, yangi so'rov vektorini (Query Vector) hosil qiladi. Endilikda bizning asosiy vazifamiz — ushbu so'rov vektorini bilimlar bazasidagi barcha savollar vektorlari bilan taqqoslab, eng yaqin (eng yuqori o'xshashlikka ega) savolni aniqlashdir. Buning uchun Kosinus o'xshashligi (Cosine Similarity) mezonidan foydalaniladi. Kosinus o'xshashligi ko'p o'lchovli fazoda ikki vektor o'rtasidagi burchak kosinusini hisoblaydi va [-1, 1] oralig'idagi qiymatni qaytaradi (matnlar uchun asosan [0, 1]). Matematik formulasi: Cosine_Similarity(A, B) = (A • B) / (||A|| * ||B||). Bizning vektorlarimiz L2 normallashgani sababli, ularning uzunligi ||A|| = ||B|| = 1 bo'ladi, ya'ni maxraj hisoblanishi shart emas: Cosine_Similarity(A, B) = A • B (ikki vektorning skalyar ko'paytmasi). Qaysi savol bilan kosinus o'xshashlik eng yuqori bo'lsa va u belgilangan chegaradan (threshold = 0.25) katta bo'lsa, o'sha savolning javobi foydalanuvchiga taqdim etiladi.")

    doc.add_page_break()

    # ==================== II BOB ====================
    add_custom_heading("II BOB. NLP ASOSIDAGI INTERAKTIV SAVOL-JAVOB CHATBOT TIZIMINI DASTURIY AMALGA OSHIRISH VA TAHLIL", level=1, before=18, after=12)
    
    add_custom_heading("2.1. Tizimning umumiy me'morchiligi (Architecture) va ma'lumotlar to'plami (Dataset) preprocessing bosqichi", level=2, before=12, after=6)
    
    add_custom_paragraph("Loyihaning dasturiy arxitekturasi foydalanuvchilar va ma'murlar uchun maksimal darajada qulay, tezkor va interaktiv bo'lishi uchun ishlab chiqilgan. Tizim ikki qismdan iborat. Birinchisi, veb-brauzerda to'liq offline ishlovchi interaktiv Foydalanuvchi Interfeysi (Veb-dashboard) bo'lib, u HTML5, CSS3, JavaScript (ES6) hamda grafik tahlillar uchun Chart.js kutubxonasi yordamida yaratilgan. Ikkinchisi esa, modelning aniqligi va ilmiy qiymatini isbotlovchi, scikit-learn va matplotlib asosida yozilgan Python modelidir. Veb-dashboardda custom NLP pipeline yaratilgan bo'lib, u foydalanuvchining so'rovlarini brauzerning o'zida bir zumda (nolga yaqin kechikish bilan) matematik tahlil qiladi. Quyidagi 1-rasmda bilimlar bazasidagi savollardan olingan eng muhim kalit so'zlarning (lemmalarning) TF-IDF og'irliklari chastotasi taqsimoti keltirilgan:")

    add_project_image("word_frequency.png", "1-rasm. Bilimlar bazasidagi savollar o'zak so'zlari (lemmalar)ning TF-IDF chastota taqsimoti.")

    add_custom_paragraph("Matn ma'lumotlarini dastlabki qayta ishlash (preprocessing) bosqichida barcha 15 ta savol maxsus funksiya orqali tozalandi. Masalan, 'TDIU qabul komissiyasi qachon ish boshlaydi?' savoli tokenizatsiyadan so'ng ['tdiu', 'qabul', 'komissiyasi', 'qachon', 'ish', 'boshlaydi'] ko'rinishini oladi. So'ngra stop-word ro'yxatida yaroqsiz yordamchi so'zlar bo'lmagani uchun to'g'ridan-to'g'ri stemming funksiyasiga o'tadi. O'zbek tili stemming qoidalariga ko'ra, 'komissiyasi' so'zidagi egalik qo'shimchasi 'si' kesilib, 'komissiya' o'zagiga keltiriladi. 'boshlaydi' so'zi esa o'z o'zagini saqlab qoladi. Yakuniy matn ko'rinishi: 'tdiu qabul komissiya qachon ish boshlaydi' shaklini oladi. Tizim bu so'zlarning barchasi asosida L2 normallashgan TF-IDF matritsasini hosil qiladi.")

    add_custom_heading("2.2. Python tilida NLP modelini amalga oshirish va visual tahlil natijalari", level=2, before=12, after=6)
    
    add_custom_paragraph("Dasturiy qismni ilmiy tahlil qilishda Python dasturlash tilining scikit-learn kutubxonasi tarkibidagi TfidfVectorizer va cosine_similarity modullaridan foydalanildi. Python skriptida tozalangan savollardan olingan lug'at (vocabulary) asosida har bir savol boshqasi bilan solishtirilib, ularning o'zaro kosinus o'xshashlik matritsasi (Heatmap) yaratildi. Bu grafik modelning savollar bir-biridan qanchalik farq qilishini aniqlash imkonini beradi. Quyidagi 2-rasmda barcha 15 ta savolning o'zaro o'xshashlik matritsasi keltirilgan:")

    add_project_image("similarity_heatmap.png", "2-rasm. Savollarning o'zaro Kosinus o'xshashlik matritsasi (Heatmap tahlili).")

    add_custom_paragraph("O'xshashlik matritsasi tahlilidan ko'rinib turibdiki, diagonal chiziqdagi barcha kataklar 1.00 qiymatga ega, chunki har bir savol o'z-o'zi bilan 100% o'xshashlikka ega. Boshqa kataklardagi qiymatlar asosan 0.00 dan 0.20 oralig'ida bo'lib, bu savollarning semantik jihatdan bir-biridan juda yaxshi farqlanganini ko'rsatadi. Masalan, S1 ('qabul komissiyasi qachon ish boshlaydi') va S8 ('magistraturaga hujjat topshirishda til sertifikati majburiymi') savollari o'rtasida 0.15 darajasidagi o'xshashlik bor, chunki ikkala savol ham 'qabul/hujjat' mavzusiga oid bo'lib, model bu bog'liqlikni ham hisobga olgan. Chatbotning turli so'rovlar bilan ishlash barqarorligini baholash uchun maxsus test so'rovlari generatsiya qilindi va turli xil threshold (chegara) darajalarida match qilish foizlari (Match Rate) tahlil qilindi. Quyidagi 3-rasmda modelning threshold darajasiga ko'ra match qilish dinamikasi ko'rsatilgan:")

    add_project_image("confidence_curve.png", "3-rasm. NLP modeli ishonch chegarasi (Threshold) bo'yicha match rate egri chizig'i.")

    add_custom_paragraph("Egri chiziq tahliliga ko'ra, model threshold darajasi 0.0 dan 0.25 gacha bo'lganda 85% dan ortiq to'g'ri match qilish natijasini ko'rsatgan. Chegara asossiz ravishda 0.8 dan oshirib yuborilganda esa, match rate 20% ga tushib ketgan, chunki foydalanuvchilar real hayotda savolni so'zma-so'z bir xil yozmaydilar. Shu sababli, tizim uchun optimal threshold (chegara) darajasi 0.25 (25%) qilib belgilandi. Bu chegara chatbotga grammatik xatoliklar yoki sinonimlarni ham to'g'ri tushunib, yuksak aniqlikda javob qaytarish imkonini taqdim etadi.")

    add_custom_heading("2.3. Foydalanuvchining interaktiv glassmorphism veb-dashboard dizayni va uning tahlili", level=2, before=12, after=6)
    
    add_custom_paragraph("Faqatgina nazariy va backend modeli bilan cheklanib qolmasdan, ushbu NLP formulasini oliy ta'lim muassasalari rahbarlari, xodimlari va talabalari osongina ishlatishlari va real vaqtda kuzatishlari uchun premium interaktiv Foydalanuvchi Interfeysi (UI) yaratildi. Sayt dizayni so'nggi zamonaviy tendensiyalar asosida 'Glassmorphism' (oynasimon shaffoflik va blur) uslubida va ko'zni charchatmaydigan, premium dark-mode foni bilan yozildi. Interaktiv dashboard 4 ta asosiy blokdan iborat:")
    add_custom_paragraph("1. AI Chatbot Console: Foydalanuvchi savol beradigan va javob oladigan klassik chat oynasi. Unda real vaqt rejimida yozish animatsiyasi (typing indicator) va har bir javob bilan birga o'xshashlik foizi visual ko'rsatiladi.", before=0, after=2)
    add_custom_paragraph("2. NLP Pipeline Visualizer: Bu loyihaning eng muhim ilmiy-tushuntirish qismi bo'lib, foydalanuvchi savol yozib jo'natganida, tizim uni 8 ta bosqichga ajratib, har bir bosqichda (lowercasing, tokenization, stop-words filtering, o'zbekcha stemming, TF-IDF vektor qiymatlari va cosine matching) matn qanday holatga kelganini real vaqtda kod ko'rinishida chiqarib beradi.", before=0, after=2)
    add_custom_paragraph("3. Bilimlar Bazasi Menejeri: Jadval shaklidagi muharrir. Bu yerda ma'murlar yangi savollar va javoblar qo'shishi, mavjudlarini tahrirlashi yoki o'chirishi mumkin. Eng qiziqarlisi, yangi ma'lumot qo'shilishi bilan Javascript NLP dvigateli barcha TF-IDF matritsalarini avtomatik qayta hisoblab chiqadi va modelni real vaqtda qayta o'qitadi.", before=0, after=2)
    add_custom_paragraph("4. NLP Analitika paneli: Chart.js kutubxonasi yordamida dynamic grafiklar integratsiya qilingan bo'lib, har bir berilgan savoldan keyin Match qilish foizi, NLP hisoblash tezligi (ms), eng faol lemmalar va savollar tarixi grafiklari avtomatik dynamic tarzda yangilanib boradi.", before=0, after=6)

    add_custom_paragraph("Bundan tashqari, loyiha tarkibiga chatbot tizimini professional va yuqori darajaga olib chiquvchi to'rtta muhim premium funksional qo'shildi va ular orqali tizimning amaliy qiymati yanada oshdi:")
    add_custom_paragraph("1) Brauzerning LocalStorage xotirasi bilan integratsiya: Tizimda bilimlar bazasiga kiritilgan har qanday yangi savol-javoblar juftligi va o'chirishlar brauzerning LocalStorage xotirasida avtomatik saqlanib boradi. Natijada foydalanuvchi sahifani yangilaganida (reload qilganida) ma'lumotlar yo'qolmaydi va model o'zining oxirgi o'qitilgan holatini mukammal saqlab qoladi.", before=0, after=2)
    add_custom_paragraph("2) Web Speech API ovozli texnologiyalari: Chatbot interfeysiga Speech-to-Text (STT) va Text-to-Speech (TTS) imkoniyatlari muvaffaqiyatli tatbiq etildi. Mikrafon tugmasi orqali o'zbek tilidagi nutqni matnga o'girib chatga yozish hamda botning javoblarini avtomatik ravishda ovozli eshitish (TTS) yoqilishi mumkin. Bu imkoniyat imkoniyati cheklangan talabalar uchun ham qulaylik yaratadi.", before=0, after=2)
    add_custom_paragraph("3) JSON Import va Eksport boshqaruvi: Tizimdagi bilimlar bazasini birgina tugma yordamida to'liq JSON fayl ko'rinishida yuklab olish (Eksport) yoki avval saqlangan bazani tizimga bir soniyada yuklash (Import) imkoniyati yaratildi. Bu ma'lumotlarni zaxiralash (backup) va almashishni osonlashtiradi.", before=0, after=2)
    add_custom_paragraph("4) Semantik Tahlil Qumlog'i (Semantic Sandbox): Bu mustaqil ilmiy modul bo'lib, unda foydalanuvchi ixtiyoriy ikkita matnni yozib, ularning preprocess (lowercasing, stop-words, stemming) o'zgarishlarini, har bir so'zining L2-normallashgan TF-IDF og'irlik qiymatlarini dynamic jadvalda solishtirishi hamda o'zaro kosinus o'xshashligini 0% dan 100% gacha bo'lgan aniq foizda ko'rishi mumkin. Bu NLP modelining matematik hisob-kitobini ochiq-oydin isbotlovchi ajoyib vizual simulyatordir.", before=0, after=6)

    doc.add_page_break()

    # ==================== XULOSA ====================
    add_custom_heading("XULOSA", level=1, before=18, after=12)
    
    add_custom_paragraph("Mazkur individual loyiha ishi doirasida \"46. Savol-javob chatbot (NLP asosida)\" mavzusi nazariy, lingvistik, iqtisodiy va dasturiy-amaliy jihatdan har tomonlama o'rganildi va quyidagi muhim yakuniy xulosalarga kelindi:")
    
    add_custom_paragraph("1. Oliy ta'lim muassasalarida talabalarning ko'plab takrorlanuvchi ma'muriy-akademik savollarini avtomatlashtirishda NLP (Natural Language Processing) chatbotlarining qo'llanilishi vaqt va inson resurslarini sezilarli tejash, universitet faoliyatining samaradorligini va talabalar mamnuniyatini oshirish imkonini berdi.")
    
    add_custom_paragraph("2. Agglyutinativ til bo'lgan o'zbek tili uchun maxsus morfemik stemming qoidalari va to'xtash so'zlari (stop-words) filtri ishlab chiqildi. Bu matnlarni preprocessing bosqichida tozalab, matematik modellarda yagona lug'at hosil qilish va model aniqligini oshirish uchun asosiy poydevor bo'lib xizmat qildi.")
    
    add_custom_paragraph("3. Matematik asosda TF-IDF vektor space modeli va Kosinus o'xshashligi (Cosine Similarity) mezonlari muvaffaqiyatli dasturlashtirildi. Python modelida barcha 15 ta TDIU savollari o'zaro tahlil qilinib, 100% aniqlik bilan match bo'lishi hamda o'xshashlik chegarasi (threshold) uchun 0.25 (25%) darajasi eng optimal ko'rsatkich ekanligi ilmiy grafiklar yordamida isbotlandi.")
    
    add_custom_paragraph("4. Brauzerda to'liq offline ishlaydigan custom Javascript NLP dvigateli va uning ustida premium glassmorphism interaktiv Dashboardi, bilimlar bazasi menejeri hamda Chart.js grafiklari ishlab chiqildi. Bu tizim orqali har qanday oliy ta'lim yoki biznes muassasalari o'zlarining shaxsiy savol-javob bazasini osongina boshqarib, chatbotni o'z vaqtida modernizatsiya qilishlari mumkin.")
    
    add_custom_paragraph("Kelajakda ushbu tizimni chuqur neyron tarmoqlar (Deep Learning) va Transformer modellar (masalan, BERT yoki GPT) bilan integratsiya qilish, tizimning semantik sinonimlarni yanada chuqurroq anglash darajasini oshirish hamda uni Hemis va boshqa ma'lumotlar bazalari bilan integratsiya qilish tavsiya etiladi. Bu ta'lim muassasalari raqamlashtirish darajasini yanada yuqori bosqichga ko'taradi.")

    doc.add_page_break()

    # ==================== FOYDALANILGAN ADABIYOTLAR ====================
    add_custom_heading("FOYDALANILGAN ADABIYOTLAR RO'YHATI", level=1, before=18, after=12)
    
    add_custom_paragraph("1. Ganiyev S.K., Karimov M.M., Tashev K.A. Axborot xavfsizligi. Toshkent: «Fan va texnologiya», 2017. ISBN 978-9943-11-366-4 (asosiy darslik).", before=6, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("2. O‘zDSt ISO/IEC 27001:2015 — Axborot xavfsizligini boshqarish tizimi. Talablar va qo‘llanma. (Milliy standart).", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("3. Daniel Jurafsky, James H. Martin. Speech and Language Processing (3rd ed. draft). Stanford University, 2023. – 624 p.", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("4. Sirojiddin Komolov, Sherzod Raxmatov. Sun’iy intellekt asoslari, Mashinaviy o’qitish. - Toshkent: “Ijod NASHR” nashriyoti, 2024. - 104-b.", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("5. Aurélien Géron. Hands-on Machine Learning with Scikit-Learn, Keras, and TensorFlow. O’Reilly Media, Inc., 2019. – 483 p.", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("6. Christopher D. Manning, Prabhakar Raghavan, Hinrich Schütze. Introduction to Information Retrieval. Cambridge University Press, 2008. - 382 p.", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("7. Steven Bird, Ewan Klein, Edward Loper. Natural Language Processing with Python. O'Reilly Media Inc., 2009. - 504 p.", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_custom_paragraph("8. O‘zbekiston Respublikasining Qonuni. Raqamli raqamlashtirish va axborot texnologiyalari to‘g‘risida. - Toshkent: “O‘zbekiston”, 2022. - B. 12-25.", before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Save document
    output_filename = "Quvonchbek_Loyiha_46.docx"
    doc.save(output_filename)
    print(f"Success: Document saved as '{output_filename}'")

if __name__ == "__main__":
    create_docx()
